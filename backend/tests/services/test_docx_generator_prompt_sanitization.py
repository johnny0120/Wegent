# SPDX-FileCopyrightText: 2025 Weibo, Inc.
#
# SPDX-License-Identifier: Apache-2.0

"""
Unit tests for docx_generator prompt sanitization.

Verifies that system-injected metadata blocks (<system-reminder>) stored in
Subtask.prompt are stripped before being written to the exported DOCX document.
"""

import json
from datetime import datetime
from unittest.mock import MagicMock

import pytest
from docx import Document

from app.services.export.docx_generator import (
    _add_file_attachment,
    _add_message,
    _sanitize_xml_text,
    generate_task_docx,
)


def _make_user_subtask(prompt: str, contexts=None):
    """Create a minimal mock Subtask with USER role."""
    subtask = MagicMock()
    subtask.role = MagicMock()
    subtask.role.value = "USER"
    subtask.prompt = prompt
    subtask.result = None
    subtask.contexts = contexts or []
    subtask.sender_user_id = None
    subtask.updated_at = datetime(2025, 1, 1, 12, 0, 0)
    return subtask


def _make_assistant_subtask(result_value: str):
    """Create a minimal mock Subtask with ASSISTANT role."""
    subtask = MagicMock()
    subtask.role = MagicMock()
    subtask.role.value = "ASSISTANT"
    subtask.prompt = None
    subtask.result = {"value": result_value}
    subtask.contexts = []
    subtask.sender_user_id = None
    subtask.updated_at = datetime(2025, 1, 1, 12, 0, 0)
    return subtask


def _make_task_and_user():
    """Create mock task (Kind) and user objects."""
    task = MagicMock()
    task.user_id = 1
    task.json = {"spec": {"teamRef": {"name": "TestTeam"}}}
    user = MagicMock()
    user.user_name = "Alice"
    return task, user


class TestDocxGeneratorPromptSanitization:
    """Tests for prompt sanitization in _add_message."""

    def test_removes_only_xml_10_forbidden_characters(self) -> None:
        """XML sanitization preserves valid Unicode and whitespace."""
        valid_text = "中文😀\t\n\r\u007f\u0085"
        text = f"A\x00\x08\x0b\x0c\x1f{valid_text}\ud800\ufffeB"

        assert _sanitize_xml_text(text) == f"A{valid_text}B"

    def test_xml_control_character_removed_from_prompt(self) -> None:
        """An XML-forbidden separator in a prompt must not break DOCX export."""
        doc = Document()
        task, user = _make_task_and_user()
        subtask = _make_user_subtask(
            "必须包含2\u001e3处合理的业务小缺陷，其他内容保持不变。"
        )

        db = MagicMock()
        db.query.return_value.filter.return_value.first.return_value = user

        _add_message(doc, subtask, task, user, db)

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "必须包含23处合理的业务小缺陷，其他内容保持不变。" in full_text
        assert "\u001e" not in full_text

    def test_xml_control_character_removed_from_task_title(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """An XML-forbidden character in the task title is removed."""
        task = MagicMock()
        task.id = 1
        task.user_id = 1
        task.json = {"metadata": {"name": "Task\x00Title"}, "spec": {}}
        db = MagicMock()
        monkeypatch.setattr(
            "app.services.export.docx_generator.subtask_store.list_by_task_ordered",
            lambda *args, **kwargs: [],
        )

        buffer = generate_task_docx(task, db)

        generated_doc = Document(buffer)
        assert generated_doc.core_properties.title == "TaskTitle"
        assert "TaskTitle" in "\n".join(p.text for p in generated_doc.paragraphs)

    def test_xml_control_character_removed_from_attachment_name(self) -> None:
        """An XML-forbidden character in an attachment name is removed."""
        doc = Document()
        attachment = MagicMock()
        attachment.name = "report\x00.docx"
        attachment.type_data = {
            "file_extension": ".docx",
            "file_size": 1024,
        }

        _add_file_attachment(doc, attachment)

        assert "report.docx" in "\n".join(p.text for p in doc.paragraphs)

    def test_plain_text_prompt_written_to_document(self):
        """Plain text user prompt appears as-is in DOCX content."""
        doc = Document()
        task, user = _make_task_and_user()
        subtask = _make_user_subtask("What is the capital of France?")

        db = MagicMock()
        db.query.return_value.filter.return_value.first.return_value = user

        _add_message(doc, subtask, task, user, db)

        # Collect all paragraph text from the document
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "What is the capital of France?" in full_text

    def test_knowledge_base_context_not_written_to_document(self) -> None:
        """Knowledge base runtime metadata must not appear in DOCX exports."""
        knowledge_base = MagicMock()
        knowledge_base.context_type = "knowledge_base"
        knowledge_base.name = "Private Knowledge Base"
        knowledge_base.type_data = {"document_count": 12}

        doc = Document()
        task, user = _make_task_and_user()
        subtask = _make_user_subtask(
            "Answer from available knowledge", [knowledge_base]
        )

        db = MagicMock()
        db.query.return_value.filter.return_value.first.return_value = user

        _add_message(doc, subtask, task, user, db)

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Answer from available knowledge" in full_text
        assert "Private Knowledge Base" not in full_text
        assert "[KB]" not in full_text

    def test_system_reminder_stripped_from_docx_output(self):
        """<system-reminder> block in JSON array prompt must NOT appear in DOCX."""
        raw_prompt = json.dumps(
            [
                {"type": "text", "text": "Summarize the attached report"},
                {
                    "type": "text",
                    "text": "<system-reminder><CurrentTime>2025-01-01 09:00</CurrentTime></system-reminder>",
                },
            ]
        )
        doc = Document()
        task, user = _make_task_and_user()
        subtask = _make_user_subtask(raw_prompt)

        db = MagicMock()
        db.query.return_value.filter.return_value.first.return_value = user

        _add_message(doc, subtask, task, user, db)

        full_text = "\n".join(p.text for p in doc.paragraphs)

        # User text must be present
        assert "Summarize the attached report" in full_text
        # System metadata must NOT appear
        assert "<system-reminder>" not in full_text
        assert "CurrentTime" not in full_text
        # Raw JSON array bracket must NOT appear as literal text
        assert full_text.count("[{") == 0

    def test_json_array_with_multiple_system_blocks(self):
        """All extra system blocks are removed; only user text is written."""
        raw_prompt = json.dumps(
            [
                {"type": "text", "text": "Review this codebase"},
                {
                    "type": "text",
                    "text": "<system-reminder><Attachment>repo.zip</Attachment></system-reminder>",
                },
                {
                    "type": "text",
                    "text": "<system-reminder><CurrentTime>2025-03-01</CurrentTime></system-reminder>",
                },
            ]
        )
        doc = Document()
        task, user = _make_task_and_user()
        subtask = _make_user_subtask(raw_prompt)

        db = MagicMock()
        db.query.return_value.filter.return_value.first.return_value = user

        _add_message(doc, subtask, task, user, db)

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Review this codebase" in full_text
        assert "<system-reminder>" not in full_text

    def test_assistant_message_uses_result_not_prompt(self):
        """ASSISTANT messages render result.value, not prompt."""
        doc = Document()
        task, user = _make_task_and_user()
        subtask = _make_assistant_subtask("Paris is the capital of France.")

        db = MagicMock()
        db.query.return_value.filter.return_value.first.return_value = user

        _add_message(doc, subtask, task, user, db)

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Paris is the capital of France." in full_text
